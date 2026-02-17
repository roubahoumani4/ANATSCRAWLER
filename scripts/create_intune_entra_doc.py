from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt


def add_heading(doc, text, level=1):
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    run.font.size = Pt(14 if level==1 else 12)


def add_paragraph(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p


def add_table(doc, headers, rows, style='Table Grid'):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = style
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
    for row in rows:
        r_cells = table.add_row().cells
        for i, cell in enumerate(row):
            r_cells[i].text = str(cell)
    return table


def build_document(path):
    doc = Document()
    doc.core_properties.title = 'Intune & Entra: Applications Design, Assessment and Implementation'
    doc.core_properties.author = 'ANAT Security'

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = title.add_run('Intune & Entra Applications — Assessment and Implementation')
    r.bold = True
    r.font.size = Pt(16)

    doc.add_paragraph()

    # Part 1: Enterprise apps in Entra
    add_heading(doc, 'Part 1 — Entra Enterprise Applications', level=1)
    add_paragraph(doc, 'Scope: Applications accessed through a web browser and integrated via Entra for SSO and access management.')

    add_heading(doc, 'Assessment and Recommendation Process', level=2)
    add_paragraph(doc, 'Before creating any application entry in Entra, perform an assessment and document recommendations. The recommendation document should include the following elements:')

    rec_headers = ['Item','Description / Guidelines']
    rec_rows = [
        ('Application name','Official vendor name and internal reference'),
        ('Purpose','Business purpose and data classification'),
        ('Authentication','Supported protocols (SAML/OAuth/OIDC) and MFA requirements'),
        ('User Mapping','Which groups/users should have access'),
        ('Provisioning','Provisioning support (SCIM/etc.)'),
        ('Risk','Data exposure and third-party risk assessment'),
        ('Recommendation','Approve/Block/Conditional access rules and next steps')
    ]
    add_table(doc, rec_headers, rec_rows)

    add_heading(doc, 'Entra Application Assignment and Lifecycle', level=2)
    add_paragraph(doc, 'After the app is created in Entra, assign it to users and groups to enable SSO. Monitor account status and revoke access or remove registrations when users leave the company (for example, remove registered sessions or block access to third-party services such as ChatGPT).')

    add_heading(doc, 'Entra — Sample Application Record (Web SSO)', level=3)
    app_headers = ['Field','Value']
    app_rows = [
        ('Name','Example WebApp'),
        ('Sign-on method','SAML 2.0'),
        ('Reply URL','https://app.example.com/sso/callback'),
        ('Identifier (Entity ID)','https://app.example.com'),
        ('User assignment','Group: ANAT-SEC-G1'),
        ('Conditional Access','MFA required, compliant device for admin tasks'),
    ]
    add_table(doc, app_headers, app_rows)

    # Part 2: Intune applications
    add_heading(doc, 'Part 2 — Applications in Intune', level=1)
    add_paragraph(doc, 'Intune delivers apps to devices based on group assignments. This section decomposes packaging and assignment strategies into two categories: wrapped (Win32) apps and built-in Microsoft Store apps.')

    # Wrapped apps strategy
    add_heading(doc, '2.1 Wrapped (Win32) Applications', level=2)
    add_paragraph(doc, 'Strategy and prerequisites for wrapping a Win32 application before uploading to Intune:')
    steps = [
        'The user performing packaging must be an Intune administrator with access to the Intune dashboard.',
        'On the packager machine: install Git and clone the Microsoft Win32 Content Prep Tool repository.',
        'Download the application installer (EXE/MSI) to the packager machine.',
        'Run the IntuneWinAppUtil.exe and provide: path to the installer, installer file name, and the output folder for the generated .intunewin file.',
        'In the Microsoft Endpoint Manager (Intune): Apps -> All apps -> Create -> Windows app (Win32). Upload the generated .intunewin file and configure application metadata, requirements, detection rules, return codes, and assignments.'
    ]
    for s in steps:
        add_paragraph(doc, f'- {s}')

    add_heading(doc, 'Packaging Example — Notepad++ (wrapped)', level=3)
    add_paragraph(doc, 'Below is a representative configuration used in Intune for a wrapped Win32 application:')

    info_headers = ['App Information','Configured Value']
    info_rows = [
        ('Name','npp.8.9.Installer.x64.exe'),
        ('Publisher','Notepad++'),
        ('Install command','npp.8.9.Installer.x64.exe /S'),
        ('Uninstall command','"C:\\Program Files\\Notepad++\\uninstall.exe" /S'),
        ('Install behavior','System'),
        ('OS minimum','Windows 10 1607'),
        ('Detection rule','File exists: C:\\Program Files\\Notepad++\\'),
    ]
    add_table(doc, info_headers, info_rows)

    assign_headers = ['Assignment Setting','Value (example)']
    assign_rows = [
        ('Group mode','Group'),
        ('Included group','ANAT-SEC-G1'),
        ('Availability','Required — devices in group will install'),
        ('End user notifications','Show all toast notifications'),
    ]
    add_table(doc, assign_headers, assign_rows)

    # Built-in apps
    add_heading(doc, '2.2 Built-in / Microsoft Store Applications', level=2)
    add_paragraph(doc, 'Built-in or Microsoft Store apps are added via Intune using the Windows Store for Business or the built-in app type. Steps:')
    built_steps = [
        'In Intune: Apps -> Add -> Choose Microsoft Store app or built-in app type.',
        'Search or specify the app and configure assignments and requirements.',
        'Assign by user or device groups; Intune will install for devices/users matching the assignment.'
    ]
    for s in built_steps:
        add_paragraph(doc, f'- {s}')

    # Appendix: Assessment checklist and Recommendation template
    add_heading(doc, 'Appendix — Assessment Checklist', level=2)
    checklist_headers = ['Check','Yes/No','Notes']
    checklist_rows = [
        ('App purpose and data classification','', ''),
        ('Auth protocol supported (SAML/OAuth)','',''),
        ('MFA required','', ''),
        ('Provisioning supported','', ''),
        ('Third-party risk acceptable','', ''),
    ]
    add_table(doc, checklist_headers, checklist_rows)

    add_heading(doc, 'Recommendation Document Template', level=2)
    add_paragraph(doc, 'Use the template below to present the final recommendation to stakeholders:')
    rec_t_headers = ['Section','Content to include']
    rec_t_rows = [
        ('Summary','Short description and recommendation (Approve / Conditional / Reject)'),
        ('Risks','Identified risks and mitigation steps'),
        ('Access','Who should be granted access and removal policy'),
        ('Implementation plan','Packaging, testing, rollout schedule'),
        ('Rollback plan','Steps to revert changes if needed')
    ]
    add_table(doc, rec_t_headers, rec_t_rows)

    # Save
    doc.save(path)


if __name__ == '__main__':
    import os
    out = os.path.join(os.getcwd(), 'Intune_Entra_Applications.docx')
    build_document(out)
    print('Generated', out)
